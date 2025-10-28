import streamlit as st
import ollama
from PyPDF2 import PdfReader
import docx

st.set_page_config(page_title="🦙 GenAI Lab — Local Llama3", layout="wide")

st.title("🦙 GenAI Lab — Powered by Local Llama3")
st.caption("Chat • Summarize • Analyze • Q&A — all fully offline using Ollama")

tab1, tab2, tab3, tab4 = st.tabs(["💬 Chat Assistant", "📄 Summarizer", "🧠 Semantic Analyzer", "🔍 Doc Q&A"])

# ========== 1️⃣ CHAT ASSISTANT ==========
with tab1:
    st.header("💬 Chat Assistant")
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []

    user_input = st.text_area("Enter your question:", key="chat_input", height=120)

    if st.button("Ask", key="chat_button"):
        if user_input.strip():
            st.session_state.chat_history.append(("You", user_input))
            with st.spinner("Thinking..."):
                response = ollama.chat(model="llama3", messages=[{"role": "user", "content": user_input}])
            reply = response['message']['content']
            st.session_state.chat_history.append(("Llama3", reply))

    for role, msg in st.session_state.chat_history:
        if role == "You":
            st.markdown(f"**🧑‍💻 You:** {msg}")
        else:
            st.markdown(f"**🤖 Llama3:** {msg}")


# ========== 2️⃣ SUMMARIZER ==========
with tab2:
    st.header("📄 Summarizer (PDF / DOC / Text Input)")
    uploaded_file = st.file_uploader("Upload PDF or Word document", type=["pdf", "docx"])
    text_input = st.text_area("Or paste your text below:", height=150, key="summ_text")

    doc_text = ""

    if uploaded_file:
        ext = uploaded_file.name.split(".")[-1].lower()
        if ext == "pdf":
            pdf_reader = PdfReader(uploaded_file)
            doc_text = "\n".join([page.extract_text() for page in pdf_reader.pages if page.extract_text()])
        elif ext == "docx":
            doc = docx.Document(uploaded_file)
            doc_text = "\n".join([p.text for p in doc.paragraphs])

    final_text = text_input.strip() if text_input.strip() else doc_text

    if st.button("Summarize", key="summarize_button"):
        if final_text:
            prompt = f"Summarize this document concisely:\n\n{final_text[:4000]}"
            with st.spinner("Summarizing..."):
                response = ollama.chat(model="llama3", messages=[{"role": "user", "content": prompt}])
            st.subheader("🧾 Summary:")
            st.write(response["message"]["content"])
        else:
            st.warning("Please upload a file or enter some text.")


# ========== 3️⃣ SEMANTIC ANALYZER ==========
with tab3:
    st.header("🧠 Semantic Analyzer (Meaning, Tone, and Emotion)")
    uploaded_sem_file = st.file_uploader("Upload PDF or DOC", type=["pdf", "docx"], key="sem_file")
    sem_text = st.text_area("Or paste text to analyze:", height=150, key="sem_text")

    sem_doc = ""
    if uploaded_sem_file:
        ext = uploaded_sem_file.name.split(".")[-1].lower()
        if ext == "pdf":
            reader = PdfReader(uploaded_sem_file)
            sem_doc = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        elif ext == "docx":
            d = docx.Document(uploaded_sem_file)
            sem_doc = "\n".join([p.text for p in d.paragraphs])

    final_sem_text = sem_text.strip() if sem_text.strip() else sem_doc

    if st.button("Analyze", key="analyze_button"):
        if final_sem_text:
            prompt = (
                "Perform semantic analysis of the following text. "
                "Explain tone, sentiment, emotion, key themes, and intent:\n\n"
                f"{final_sem_text[:4000]}"
            )
            with st.spinner("Analyzing..."):
                response = ollama.chat(model="llama3", messages=[{"role": "user", "content": prompt}])
            st.subheader("🔍 Semantic Analysis Result:")
            st.write(response["message"]["content"])
        else:
            st.warning("Please upload a file or enter text to analyze.")


# ========== 4️⃣ DOC Q&A ==========
with tab4:
    st.header("🔍 Ask Questions from Document (PDF / DOC / Text)")
    uploaded_qa_file = st.file_uploader("Upload document", type=["pdf", "docx"], key="qa_file")
    question = st.text_input("Enter your question:", key="qa_question")

    doc_content = ""
    if uploaded_qa_file:
        ext = uploaded_qa_file.name.split(".")[-1].lower()
        if ext == "pdf":
            reader = PdfReader(uploaded_qa_file)
            doc_content = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        elif ext == "docx":
            d = docx.Document(uploaded_qa_file)
            doc_content = "\n".join([p.text for p in d.paragraphs])

    if st.button("Ask Document", key="qa_button"):
        if question.strip() and doc_content.strip():
            prompt = f"Based on the following content, answer this question clearly.\n\nContent:\n{doc_content[:4000]}\n\nQuestion: {question}"
            with st.spinner("Searching..."):
                response = ollama.chat(model="llama3", messages=[{"role": "user", "content": prompt}])
            st.subheader("📘 Answer:")
            st.write(response["message"]["content"])
        else:
            st.warning("Please upload a document and enter a question.")
